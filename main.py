import io
import json
import os
import tempfile
from fastapi import FastAPI, File, Form, HTTPException, UploadFile, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.openapi.utils import get_openapi
from PIL import Image, ImageOps

# Word Document Libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# PDF Libraries
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image as RLImage, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

app = FastAPI(
    title="Civil Site Inspection Report API",
    version="1.0.0"
)

# --- SWAGGER UI FILE UPLOAD FIX (OpenAPI 3.0.3 Patch) ---
def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    
    openapi_schema = get_openapi(
        title=app.title,
        version=app.version,
        routes=app.routes,
    )
    
    # Force OpenAPI 3.0.3 spec so Swagger UI renders file controls correctly
    openapi_schema["openapi"] = "3.0.3"
    
    for path in openapi_schema.get("paths", {}).values():
        for method in path.values():
            req_body = method.get("requestBody", {})
            content = req_body.get("content", {})
            if "multipart/form-data" in content:
                schema = content["multipart/form-data"].get("schema", {})
                props = schema.get("properties", {})
                if "files" in props:
                    props["files"] = {
                        "type": "array",
                        "items": {"type": "string", "format": "binary"}
                    }
                    
    app.openapi_schema = openapi_schema
    return app.openapi_schema

app.openapi = custom_openapi

# Enable CORS for Vercel, Localhost, and Mobile App access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def cleanup_temp_file(path: str):
    """Deletes temporary generated report files after HTTP download completes."""
    if os.path.exists(path):
        try:
            os.remove(path)
        except Exception:
            pass


# --- IMAGE PROCESSING & DUAL-CONSTRAINT ZERO-CROP SCALING MATH ---
def get_processed_image(file_bytes: bytes, rotation: int = 0) -> Image.Image:
    img = Image.open(io.BytesIO(file_bytes))
    
    # Auto-orient based on smartphone camera EXIF tags
    img = ImageOps.exif_transpose(img)
    
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    if rotation != 0:
        img = img.rotate(-rotation, expand=True)
        
    return img


def get_fitted_dimensions(img_obj: Image.Image, max_w_in: float, max_h_in: float):
    orig_w, orig_h = img_obj.size
    aspect_ratio = orig_w / orig_h

    calc_w = max_w_in
    calc_h = calc_w / aspect_ratio

    if calc_h > max_h_in:
        calc_h = max_h_in
        calc_w = calc_h * aspect_ratio

    return calc_w, calc_h


def prepare_temp_image(pil_img: Image.Image) -> str:
    temp_img_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    temp_img_path = temp_img_file.name
    temp_img_file.close()
    pil_img.save(temp_img_path, format="PNG")
    return temp_img_path


# --- WORD GENERATOR ENGINE ---
def create_docx_report(title: str, photo_items: list, photos_per_page: int, cols_per_page: int) -> str:
    doc = Document()
    temp_files = []

    try:
        for section in doc.sections:
            section.top_margin = Inches(0.35)
            section.bottom_margin = Inches(0.35)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)

        total_photos = len(photo_items)
        chunk_size = photos_per_page
        rows_per_page = (photos_per_page + cols_per_page - 1) // cols_per_page
        page_avail_w = 7.7  # 8.5" - 0.8" margins

        row_gap_pt = 100 if rows_per_page == 2 else (25 if rows_per_page == 3 else 10)

        for i in range(0, total_photos, chunk_size):
            is_first_page = (i == 0)
            has_title = bool(is_first_page and title)

            if has_title:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(title)
                run.font.name = "Calibri"
                run.font.size = Pt(16)
                run.font.bold = True
                p.paragraph_format.space_after = Pt(8)

            page_avail_h = 9.3 if has_title else 9.9
            max_cell_w = (page_avail_w / cols_per_page) - 0.15
            max_img_h = (page_avail_h / rows_per_page) - (1.25 if rows_per_page <= 2 else 0.8)

            chunk = photo_items[i:i + chunk_size]

            table = doc.add_table(rows=rows_per_page, cols=cols_per_page)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for r_idx in range(rows_per_page):
                for c_idx in range(cols_per_page):
                    item_idx = r_idx * cols_per_page + c_idx
                    cell = table.cell(r_idx, c_idx)
                    cell.width = Inches(page_avail_w / cols_per_page)

                    if item_idx < len(chunk):
                        item = chunk[item_idx]
                        img = get_processed_image(item["file_bytes"], item.get("rotation", 0))
                        w_in, h_in = get_fitted_dimensions(img, max_cell_w, max_img_h)
                        img_path = prepare_temp_image(img)
                        temp_files.append(img_path)

                        # Image Paragraph
                        img_p = cell.paragraphs[0]
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_p.paragraph_format.space_before = Pt(0)
                        img_p.paragraph_format.space_after = Pt(2)
                        run = img_p.add_run()
                        run.add_picture(img_path, width=Inches(w_in), height=Inches(h_in))

                        # Caption Paragraph
                        cap_p = cell.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_p.paragraph_format.space_before = Pt(2)
                        cap_p.paragraph_format.space_after = Pt(row_gap_pt) if r_idx < rows_per_page - 1 else Pt(2)

                        if item.get("caption"):
                            cap_run = cap_p.add_run(item["caption"])
                            cap_run.font.name = item.get("font_name", "Calibri")
                            cap_run.font.size = Pt(item.get("font_size", 10))
                            cap_run.font.bold = item.get("bold", False)
                            cap_run.font.italic = item.get("italic", False)

            if i + chunk_size < total_photos:
                doc.add_page_break()

        docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        doc.save(docx_path)
        return docx_path

    finally:
        for tf in temp_files:
            if os.path.exists(tf):
                try:
                    os.remove(tf)
                except Exception:
                    pass


# --- PURE PYTHON PDF GENERATOR ENGINE ---
def create_pdf_report(title: str, photo_items: list, photos_per_page: int, cols_per_page: int) -> str:
    pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    temp_files = []

    try:
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            leftMargin=0.4 * inch, rightMargin=0.4 * inch,
            topMargin=0.35 * inch, bottomMargin=0.35 * inch
        )

        styles = getSampleStyleSheet()
        story = []

        total_photos = len(photo_items)
        chunk_size = photos_per_page
        rows_per_page = (photos_per_page + cols_per_page - 1) // cols_per_page
        page_avail_w = 7.7  # inches

        bottom_padding_pt = 100 if rows_per_page == 2 else (25 if rows_per_page == 3 else 10)

        for i in range(0, total_photos, chunk_size):
            is_first_page = (i == 0)
            has_title = bool(is_first_page and title)

            if has_title:
                title_style = ParagraphStyle(
                    'DocTitle',
                    parent=styles['Heading1'],
                    alignment=1,
                    fontSize=16,
                    leading=20,
                    spaceAfter=8
                )
                story.append(Paragraph(title, title_style))

            page_avail_h = 9.3 if has_title else 9.9
            max_cell_w = (page_avail_w / cols_per_page) - 0.15
            max_img_h = (page_avail_h / rows_per_page) - (1.25 if rows_per_page <= 2 else 0.8)

            chunk = photo_items[i:i + chunk_size]
            table_data = []

            for r in range(rows_per_page):
                row_cells = []
                for c in range(cols_per_page):
                    idx = r * cols_per_page + c
                    if idx < len(chunk):
                        item = chunk[idx]
                        img = get_processed_image(item["file_bytes"], item.get("rotation", 0))
                        w_in, h_in = get_fitted_dimensions(img, max_cell_w, max_img_h)
                        img_path = prepare_temp_image(img)
                        temp_files.append(img_path)

                        rl_img = RLImage(img_path, width=w_in * inch, height=h_in * inch)

                        cap_text = item.get("caption", "")
                        tag_open = ""
                        tag_close = ""
                        if item.get("bold", False):
                            tag_open += "<b>"
                            tag_close = "</b>" + tag_close
                        if item.get("italic", False):
                            tag_open += "<i>"
                            tag_close = "</i>" + tag_close

                        font_family = "Helvetica"
                        if item.get("font_name") == "Times New Roman":
                            font_family = "Times-Roman"

                        font_size = item.get("font_size", 10)
                        p_style = ParagraphStyle(
                            f'CapStyle_{i}_{idx}',
                            alignment=1,
                            fontName=font_family,
                            fontSize=font_size,
                            leading=font_size + 2,
                            spaceBefore=2,
                            spaceAfter=2
                        )
                        cap_p = Paragraph(f"{tag_open}{cap_text}{tag_close}", p_style)

                        row_cells.append([rl_img, cap_p])
                    else:
                        row_cells.append("")

                table_data.append(row_cells)

            col_widths = [(page_avail_w / cols_per_page) * inch] * cols_per_page
            t = Table(table_data, colWidths=col_widths)

            table_styles = [
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]

            if rows_per_page > 1:
                for r in range(rows_per_page - 1):
                    table_styles.append(('BOTTOMPADDING', (0, r), (-1, r), bottom_padding_pt))

            t.setStyle(TableStyle(table_styles))

            story.append(t)
            if i + chunk_size < total_photos:
                story.append(PageBreak())

        doc.build(story)
        return pdf_path

    finally:
        for tf in temp_files:
            if os.path.exists(tf):
                try:
                    os.remove(tf)
                except Exception:
                    pass


# --- API ENDPOINTS ---
@app.get("/")
def health_check():
    return {"status": "online", "message": "Civil Site Inspection Report API Ready"}


@app.post("/generate-report")
async def generate_report(
    background_tasks: BackgroundTasks,
    files: list[UploadFile] = File(...),
    metadata_json: str = Form(...),
    export_format: str = Form("pdf"),
    doc_title: str = Form("SITE INSPECTION REPORT"),
    use_title: bool = Form(True),
    photos_per_page: int = Form(4),
    cols_per_page: int = Form(2),
):
    try:
        metadata = json.loads(metadata_json)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON format for metadata_json")

    if len(files) != len(metadata):
        raise HTTPException(
            status_code=400,
            detail="Mismatch between uploaded file count and metadata list length",
        )

    photo_items = []
    for idx, file_obj in enumerate(files):
        file_bytes = await file_obj.read()
        item = metadata[idx]
        item["file_bytes"] = file_bytes
        photo_items.append(item)

    title_to_use = doc_title if use_title else ""

    if export_format.lower() == "docx":
        out_path = create_docx_report(title_to_use, photo_items, photos_per_page, cols_per_page)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "Site_Inspection_Report.docx"
    else:
        out_path = create_pdf_report(title_to_use, photo_items, photos_per_page, cols_per_page)
        media_type = "application/pdf"
        filename = "Site_Inspection_Report.pdf"

    background_tasks.add_task(cleanup_temp_file, out_path)

    return FileResponse(out_path, media_type=media_type, filename=filename)
